"""
Convert DOCUMENTATION.md into a formatted Word (.docx) document.

Handles: headings (#..####), paragraphs, bullet lists, bold (**text**),
inline code (`code`), fenced code blocks (``` ```), and Markdown tables.
"""

import os
import re

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

HERE = os.path.dirname(os.path.abspath(__file__))
SRC = os.path.join(HERE, "DOCUMENTATION.md")
OUT = os.path.join(HERE, "DOCUMENTATION.docx")

HEADING_COLORS = {
    1: RGBColor(0x1A, 0x52, 0x76),
    2: RGBColor(0x2C, 0x3E, 0x50),
    3: RGBColor(0x34, 0x49, 0x5E),
    4: RGBColor(0x5D, 0x6D, 0x7E),
}

BOLD_RE = re.compile(r"\*\*(.+?)\*\*")
CODE_RE = re.compile(r"`([^`]+?)`")


def add_rich_text(paragraph, text):
    """Add text to a paragraph, rendering **bold** and `code` spans."""
    # Split on bold and code while keeping delimiters
    tokens = re.split(r"(\*\*.+?\*\*|`[^`]+?`)", text)
    for tok in tokens:
        if not tok:
            continue
        if tok.startswith("**") and tok.endswith("**"):
            run = paragraph.add_run(tok[2:-2])
            run.bold = True
        elif tok.startswith("`") and tok.endswith("`"):
            run = paragraph.add_run(tok[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)
        else:
            paragraph.add_run(tok)


def add_code_block(doc, code_lines):
    """Add a monospaced, shaded code block."""
    para = doc.add_paragraph()
    para.paragraph_format.left_indent = Pt(12)
    run = para.add_run("\n".join(code_lines))
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)


def add_table(doc, rows):
    """Add a Markdown table (list of cell-lists) as a Word table."""
    if not rows:
        return
    cols = len(rows[0])
    table = doc.add_table(rows=0, cols=cols)
    table.style = "Light Grid Accent 1"
    for r_idx, row in enumerate(rows):
        cells = table.add_row().cells
        for c_idx in range(cols):
            text = row[c_idx] if c_idx < len(row) else ""
            cell = cells[c_idx]
            cell.text = ""
            para = cell.paragraphs[0]
            add_rich_text(para, text)
            if r_idx == 0:
                for run in para.runs:
                    run.bold = True


def parse_table_row(line):
    parts = [c.strip() for c in line.strip().strip("|").split("|")]
    return parts


def is_table_separator(line):
    return bool(re.fullmatch(r"\|?[\s:|-]+\|?", line.strip())) and "-" in line


def convert():
    with open(SRC, "r", encoding="utf-8") as f:
        lines = f.read().split("\n")

    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)

    i = 0
    n = len(lines)
    while i < n:
        line = lines[i]
        stripped = line.strip()

        # Fenced code block
        if stripped.startswith("```"):
            i += 1
            code_lines = []
            while i < n and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            add_code_block(doc, code_lines)
            i += 1  # skip closing ```
            continue

        # Table (header row followed by separator)
        if stripped.startswith("|") and i + 1 < n and is_table_separator(lines[i + 1]):
            table_rows = [parse_table_row(stripped)]
            i += 2  # skip header + separator
            while i < n and lines[i].strip().startswith("|"):
                table_rows.append(parse_table_row(lines[i]))
                i += 1
            add_table(doc, table_rows)
            doc.add_paragraph("")
            continue

        # Headings
        m = re.match(r"^(#{1,4})\s+(.*)$", stripped)
        if m:
            level = len(m.group(1))
            heading = doc.add_heading(level=level)
            run = heading.add_run(m.group(2))
            run.font.color.rgb = HEADING_COLORS.get(level, HEADING_COLORS[4])
            i += 1
            continue

        # Horizontal rule
        if stripped in ("---", "***", "___"):
            doc.add_paragraph("_" * 50).alignment = WD_ALIGN_PARAGRAPH.CENTER
            i += 1
            continue

        # Blockquote
        if stripped.startswith(">"):
            para = doc.add_paragraph(style="Intense Quote")
            add_rich_text(para, stripped.lstrip("> ").strip())
            i += 1
            continue

        # Bullet list
        bullet = re.match(r"^(\s*)[-*]\s+(.*)$", line)
        if bullet:
            indent = len(bullet.group(1))
            style = "List Bullet 2" if indent >= 2 else "List Bullet"
            para = doc.add_paragraph(style=style)
            add_rich_text(para, bullet.group(2))
            i += 1
            continue

        # Numbered list
        numbered = re.match(r"^(\s*)\d+\.\s+(.*)$", line)
        if numbered:
            para = doc.add_paragraph(style="List Number")
            add_rich_text(para, numbered.group(2))
            i += 1
            continue

        # Blank line
        if stripped == "":
            i += 1
            continue

        # Normal paragraph
        para = doc.add_paragraph()
        add_rich_text(para, stripped)
        i += 1

    doc.save(OUT)
    print(f"Saved Word document to: {OUT}")


if __name__ == "__main__":
    convert()
